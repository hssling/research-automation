import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
from prophet import Prophet
from sklearn.metrics import mean_squared_error, mean_absolute_error
from sklearn.preprocessing import MinMaxScaler
from statsmodels.tsa.arima.model import ARIMA
from tensorflow.keras import Sequential
from tensorflow.keras.layers import LSTM, Dense
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import openpyxl
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows

# ================================
# 1. Load Dataset
# ================================
df = pd.read_csv("time_series_pipeline/data/timeseries.csv")  # columns: ds (date), y (value)
df['ds'] = pd.to_datetime(df['ds'])

# ================================
# 2. Prophet Forecasting
# ================================
prophet_model = Prophet()
prophet_model.fit(df)

future = prophet_model.make_future_dataframe(periods=90)  # 90-day forecast
forecast = prophet_model.predict(future)

# Save plots
fig1 = prophet_model.plot(forecast)
fig1.savefig("time_series_pipeline/output/prophet_forecast.png")

fig2 = prophet_model.plot_components(forecast)
fig2.savefig("time_series_pipeline/output/prophet_components.png")

print("✅ Prophet forecast saved in /output")

# ================================
# 3. ARIMA Forecasting
# ================================
train, test = df['y'][:-30], df['y'][-30:]
arima_model = ARIMA(train, order=(5,1,0))
arima_fit = arima_model.fit()

forecast_arima = arima_fit.forecast(steps=30)
error = mean_squared_error(test, forecast_arima)
print(f"ARIMA MSE: {error}")

plt.figure(figsize=(10,5))
plt.plot(df['ds'][-60:], df['y'][-60:], label="Actual")
plt.plot(df['ds'][-30:], forecast_arima, label="ARIMA Forecast")
plt.legend()
plt.savefig("time_series_pipeline/output/arima_forecast.png")

print("✅ ARIMA forecast saved in /output")

# ================================
# 4. LSTM Neural Network Forecast
# ================================
values = df['y'].values.reshape(-1,1)
train_size = int(len(values)*0.8)
train, test = values[:train_size], values[train_size:]

# Normalize
scaler = MinMaxScaler()
train_scaled = scaler.fit_transform(train)
test_scaled = scaler.transform(test)

# Prepare sequences
def create_dataset(dataset, look_back=5):
    X, Y = [], []
    for i in range(len(dataset)-look_back):
        X.append(dataset[i:i+look_back,0])
        Y.append(dataset[i+look_back,0])
    return np.array(X), np.array(Y)

look_back = 5
X_train, y_train = create_dataset(train_scaled, look_back)
X_test, y_test = create_dataset(test_scaled, look_back)

X_train = X_train.reshape(X_train.shape[0], look_back, 1)
X_test = X_test.reshape(X_test.shape[0], look_back, 1)

# Build LSTM
model = Sequential([
    LSTM(50, return_sequences=False, input_shape=(look_back,1)),
    Dense(1)
])
model.compile(optimizer='adam', loss='mse')
model.fit(X_train, y_train, epochs=20, batch_size=16, verbose=0)

predictions = model.predict(X_test)
predictions = scaler.inverse_transform(predictions)

# Plot
plt.figure(figsize=(10,5))
plt.plot(df['ds'][-len(y_test):], scaler.inverse_transform(y_test.reshape(-1,1)), label="Actual")
plt.plot(df['ds'][-len(predictions):], predictions, label="LSTM Forecast")
plt.legend()
plt.savefig("time_series_pipeline/output/lstm_forecast.png")

print("✅ LSTM forecast saved in /output")

# ================================
# 5. Create Word Report
# ================================
doc = Document()
doc.add_heading("Time Series Forecasting Report", 0)

doc.add_paragraph("This report presents automated time series analysis and forecasting using Prophet, ARIMA, and LSTM models.")
doc.add_paragraph(f"Dataset length: {len(df)} observations")
doc.add_paragraph(f"ARIMA Mean Squared Error: {round(error, 2)}")

doc.add_heading("Prophet Forecast", level=1)
doc.add_picture("time_series_pipeline/output/prophet_forecast.png", width=Inches(5))
doc.add_picture("time_series_pipeline/output/prophet_components.png", width=Inches(5))

doc.add_heading("ARIMA Forecast", level=1)
doc.add_picture("time_series_pipeline/output/arima_forecast.png", width=Inches(5))

doc.add_heading("LSTM Forecast", level=1)
doc.add_picture("time_series_pipeline/output/lstm_forecast.png", width=Inches(5))

doc.add_heading("Interpretation", level=1)
doc.add_paragraph(
    "The Prophet model captures trend and seasonality effectively. "
    "ARIMA provides a strong baseline for short-term forecasting, "
    "while the LSTM model adapts to complex nonlinear patterns. "
    "Comparison of MSE suggests ARIMA is useful for stable series, "
    "while LSTM may outperform when long-term dependencies exist. "
    "Prophet is highly interpretable and well-suited for business/health forecasting."
)

doc.save("time_series_pipeline/output/report.docx")
print("✅ Word report saved in /output")

# ================================
# 6. Create PDF Report (simple)
# ================================
c = canvas.Canvas("time_series_pipeline/output/report.pdf", pagesize=letter)
c.setFont("Helvetica", 16)
c.drawString(100, 750, "Time Series Forecasting Report")
c.setFont("Helvetica", 12)
c.drawString(100, 720, f"Dataset length: {len(df)} observations")
c.drawString(100, 700, f"ARIMA MSE: {round(error, 2)}")

c.drawString(100, 670, "See detailed Word report for figures and interpretation.")
c.showPage()
c.save()

print("✅ PDF report saved in /output")

# ================================
# 7. Export Forecasts to Excel
# ================================
# First compute LSTM metrics here so we have the variables
# For LSTM, we need to align the lengths since y_test represents the full test set after sequence processing
lstm_actual_len = len(scaler.inverse_transform(y_test.reshape(-1,1)))
lstm_pred_len = len(predictions)
min_len = min(lstm_actual_len, lstm_pred_len)

lstm_actual = scaler.inverse_transform(y_test.reshape(-1,1))[:min_len]
lstm_pred = predictions[:min_len]
mse_lstm = mean_squared_error(lstm_actual, lstm_pred)
rmse_lstm = np.sqrt(mse_lstm)
mae_lstm = mean_absolute_error(lstm_actual, lstm_pred)

wb = Workbook()
ws1 = wb.active
ws1.title = "Forecasts"

# Prophet Forecast (future predictions only)
future_prophet = forecast[forecast['ds'] > df['ds'].max()]
future_prophet = future_prophet[['ds','yhat','yhat_lower','yhat_upper']]
future_prophet.columns = ['Date','Prophet_Forecast','Prophet_Lower','Prophet_Upper']

# Write future forecasts to Excel
for r in dataframe_to_rows(future_prophet, index=False, header=True):
    ws1.append(r)

# ================================
# 8. Model Comparison Sheet
# ================================
ws2 = wb.create_sheet(title="Model_Comparison")

# Compute metrics - fix the variable references
rmse_arima = np.sqrt(error)  # error is already MSE from earlier

# ARIMA MAE calculation (on the same test set)
mae_arima = mean_absolute_error(df['y'][-30:], forecast_arima)  # Use the last 30 values

comparison = pd.DataFrame({
    'Model': ['ARIMA','LSTM'],
    'MSE': [error, mse_lstm],
    'RMSE': [rmse_arima, rmse_lstm],
    'MAE': [mae_arima, mae_lstm],
    'AIC': [arima_fit.aic, None],
    'BIC': [arima_fit.bic, None]
})

# Add Prophet note
comparison.loc[len(comparison)] = ['Prophet','(see log-likelihood)','(see diagnostics)',
                                  '(not directly comparable)',None,None]

for r in dataframe_to_rows(comparison, index=False, header=True):
    ws2.append(r)

# Save workbook
wb.save("time_series_pipeline/output/forecasts.xlsx")
print("✅ Forecasts and model comparison saved in /output/forecasts.xlsx")

# ================================
# 9. Export Future Forecasts to CSV
# ================================
future_prophet.to_csv("time_series_pipeline/output/future_forecasts.csv", index=False)
print("✅ Future forecasts CSV saved in /output/future_forecasts.csv")

# ================================
# 10. Export Individual Model CSVs
# ================================

# Prophet forecast (only future predictions beyond dataset length)
prophet_future = forecast[forecast['ds'] > df['ds'].max()]
prophet_future[['ds','yhat','yhat_lower','yhat_upper']].to_csv(
    "time_series_pipeline/output/prophet_forecast.csv", index=False
)

# ARIMA forecast (next 30 days as DataFrame)
arima_df = pd.DataFrame({
    'Date': df['ds'][-30:].values if len(df) >= 30 else df['ds'].values,
    'ARIMA_Forecast': forecast_arima
})
arima_df.to_csv("time_series_pipeline/output/arima_forecast.csv", index=False)

# LSTM forecast (test set predictions)
lstm_df = pd.DataFrame({
    'Date': df['ds'][-len(predictions):].values if len(predictions) < len(df) else df['ds'][-len(predictions):].values,
    'LSTM_Forecast': predictions.flatten()
})
lstm_df.to_csv("time_series_pipeline/output/lstm_forecast.csv", index=False)

print("✅ Individual forecast CSVs saved in /output")
