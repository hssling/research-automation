#!/usr/bin/env python3
"""
Manuscript to DOCX Converter
Converts the 5 main meta-analysis manuscripts and their components into single DOCX documents
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

def setup_document_styles(doc):
    """Setup document styles for manuscript formatting"""
    # Title style
    title_style = doc.styles.add_style('ManuscriptTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Section header style
    section_style = doc.styles.add_style('ManuscriptSection', WD_STYLE_TYPE.PARAGRAPH)
    section_style.font.size = Pt(14)
    section_style.font.bold = True

    # Subsection style
    subsection_style = doc.styles.add_style('ManuscriptSubsection', WD_STYLE_TYPE.PARAGRAPH)
    subsection_style.font.size = Pt(12)
    subsection_style.font.bold = True

    # Normal text style
    normal_style = doc.styles.add_style('ManuscriptBody', WD_STYLE_TYPE.PARAGRAPH)
    normal_style.font.size = Pt(11)

def find_component_files(manuscript_name):
    """Find all component files for a given manuscript"""
    components = []

    # Special handling for PPG project
    if 'ppg' in manuscript_name.lower():
        # Hard-coded components for PPG project
        ppg_components = [
            ('protocol', 'ppg_hr_accuracy_meta_analysis/protocol.md'),
            ('search_strategy', 'ppg_hr_accuracy_meta_analysis/detailed_search_strategy.md'),
            ('data_extraction', 'ppg_hr_accuracy_meta_analysis/data_extraction_form.md'),
            ('forest_plot', 'ppg_hr_accuracy_meta_analysis/results/forest_plot_visualization.txt'),
            ('bland_altman', 'ppg_hr_accuracy_meta_analysis/results/bland_altman_plot.txt'),
            ('performance_table', 'ppg_hr_accuracy_meta_analysis/results/performance_comparison_table.md'),
            ('validation_report', 'ppg_hr_accuracy_meta_analysis/results/validation_report.md'),
            ('project_summary', 'ppg_hr_accuracy_meta_analysis/project_summary.md')
        ]
        supporting_files = [(cat, path) for cat, path in ppg_components if os.path.exists(path)]
        return f'ppg_hr_accuracy_meta_analysis/manuscript_draft.md', supporting_files

    # Define expected component patterns for other projects
    base_patterns = {
        'prisma_flow': ['prisma_flow', 'PRISMA_flow'],
        'prospero_registration': ['prospero_registration', 'PROSPERO_registration'],
        'protocol': ['protocol'],
        'appendices': ['appendices', 'supplementary', 'technical_appendices'],
        'results_tables': ['results_tables', 'results'],
        'references': ['references'],
        'validation': ['validation'],
        'plots_generator': ['plots_generator'],
        'executive_summary': ['executive_summary'],
        'supplementary_materials': ['supplementary_materials']
    }

    # Remove '_manuscript.md' from name for matching
    base_name = manuscript_name.replace('_meta_analysis_manuscript.md', '').replace('_systematic_review.md', '').replace('.md', '')

    # Special handling for microbiome allergy (different naming pattern)
    if 'microbiome' in base_name.lower():
        base_name = 'microbiome_allergy'

    # Find files containing the base name
    files_found = []
    for root, dirs, files in os.walk('.'):
        for file in files:
            if file.endswith(('.md', '.csv', '.py')):
                file_lower = file.lower()
                base_lower = base_name.lower()

                # Check if file contains the manuscript base name
                if base_lower in file_lower and 'manuscript' not in file_lower:
                    files_found.append(os.path.join(root, file))

    # Filter and categorize files
    manuscript_file = None
    supporting_files = []

    for file_path in files_found:
        file_name = os.path.basename(file_path).lower()

        # Main manuscript file
        if 'manuscript' in file_name and '.docx' not in file_name:
            manuscript_file = file_path
        else:
            # Check which category the file belongs to
            for category, patterns in base_patterns.items():
                if any(pattern.lower().replace('_', '') in file_name.replace('_', '') for pattern in patterns):
                    supporting_files.append((category, file_path))
                    break
            else:
                # Unclassified but related files
                supporting_files.append(('other', file_path))

    # Sort supporting files by category priority
    category_priority = ['protocol', 'prospero_registration', 'prisma_flow',
                        'validation', 'results_tables', 'appendices',
                        'supplementary_materials', 'references',
                        'plots_generator', 'executive_summary', 'other']

    supporting_files.sort(key=lambda x: category_priority.index(x[0]) if x[0] in category_priority else len(category_priority))

    return manuscript_file, supporting_files

def convert_markdown_to_docx(content, doc):
    """Convert markdown content to DOCX format"""
    lines = content.split('\n')
    current_list = None
    in_code_block = False
    code_content = []

    for line in lines:
        line = line.rstrip()

        if not line.strip():
            if not in_code_block:
                doc.add_paragraph('')
            continue

        # Code block handling
        if line.strip().startswith('```'):
            if in_code_block:
                # End of code block
                code_text = '\n'.join(code_content)
                p = doc.add_paragraph(code_text)
                p.style = doc.styles['ManuscriptBody']
                p.paragraph_format.left_indent = Inches(0.25)
                in_code_block = False
                code_content = []
            else:
                # Start of code block
                in_code_block = True
            continue

        if in_code_block:
            code_content.append(line)
            continue

        # Headers
        if line.startswith('# '):
            p = doc.add_paragraph(line[2:], style='ManuscriptTitle')
        elif line.startswith('## '):
            p = doc.add_paragraph(line[3:], style='ManuscriptSection')
        elif line.startswith('### '):
            p = doc.add_paragraph(line[4:], style='ManuscriptSubsection')
        elif line.startswith('#### ') or line.startswith('##### ') or line.startswith('###### '):
            p = doc.add_paragraph(line.lstrip('# '), style='ManuscriptBody')
            p.runs[0].bold = True

        # Lists
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            if current_list is None:
                current_list = doc.add_paragraph(style='ManuscriptBody')
            else:
                current_list = doc.add_paragraph(style='ManuscriptBody')
            current_list.add_run(line.strip()[2:])
            current_list.paragraph_format.left_indent = Inches(0.25)

        elif re.match(r'^\d+\.', line.strip()):
            if current_list is None:
                current_list = doc.add_paragraph(style='ManuscriptBody')
            else:
                current_list = doc.add_paragraph(style='ManuscriptBody')
            current_list.add_run(line.strip())
            current_list.paragraph_format.left_indent = Inches(0.25)

        else:
            # Regular paragraph
            p = doc.add_paragraph(line, style='ManuscriptBody')
            current_list = None

def create_combined_manuscript(manuscript_name):
    """Create a combined DOCX document for a manuscript and its components"""

    print(f"📝 Processing manuscript: {manuscript_name}")

    # Find component files
    main_file, components = find_component_files(manuscript_name)

    if not main_file:
        print(f"❌ Could not find main manuscript file for {manuscript_name}")
        return False

    print(f"📋 Found main manuscript: {os.path.basename(main_file)}")
    print(f"📋 Found {len(components)} supporting components")

    # Create new document
    doc = Document()
    setup_document_styles(doc)

    # Add title page
    title = doc.add_paragraph(f"{manuscript_name.replace('_', ' ').title()}", style='ManuscriptTitle')
    title.paragraph_format.space_after = Inches(0.5)

    subtitle = doc.add_paragraph("Complete Systematic Review and Meta-Analysis", style='ManuscriptSection')
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()

    # Table of Contents
    toc = doc.add_paragraph("Table of Contents", style='ManuscriptSection')
    toc_items = ["1. Main Manuscript", "2. Supporting Documentation"]

    for item in toc_items:
        doc.add_paragraph(item, style='ManuscriptBody')

    doc.add_page_break()

    # Main manuscript
    doc.add_paragraph("1. Main Manuscript", style='ManuscriptSection')

    try:
        with open(main_file, 'r', encoding='utf-8') as f:
            content = f.read()
        convert_markdown_to_docx(content, doc)
    except Exception as e:
        print(f"❌ Error reading main manuscript {main_file}: {e}")
        doc.add_paragraph(f"Error loading main manuscript: {str(e)}", style='ManuscriptBody')

    doc.add_page_break()

    # Supporting components
    doc.add_paragraph("2. Supporting Documentation", style='ManuscriptSection')

    section_num = 3
    for i, (category, file_path) in enumerate(components):
        if i < len(components):  # Limit to reasonable number
            doc.add_paragraph(f"{section_num}. {category.replace('_', ' ').title()}", style='ManuscriptSection')

            try:
                file_ext = os.path.splitext(file_path)[1].lower()

                if file_ext == '.md':
                    # Convert markdown
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    convert_markdown_to_docx(content, doc)

                elif file_ext == '.csv':
                    # Add CSV content as formatted text
                    doc.add_paragraph("CSV Data Content:", style='ManuscriptSubsection')
                    with open(file_path, 'r', encoding='utf-8') as f:
                        csv_content = f.read()
                    p = doc.add_paragraph(csv_content, style='ManuscriptBody')
                    p.paragraph_format.left_indent = Inches(0.25)

                elif file_ext == '.py':
                    # Add Python code as code block
                    doc.add_paragraph("Python Script Content:", style='ManuscriptSubsection')
                    with open(file_path, 'r', encoding='utf-8') as f:
                        code_content = f.read()
                    p = doc.add_paragraph(code_content, style='ManuscriptBody')
                    p.paragraph_format.left_indent = Inches(0.25)

                else:
                    # Plain text handling
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            content = f.read()
                        doc.add_paragraph(content, style='ManuscriptBody')
                    except:
                        doc.add_paragraph(f"Binary or unreadable file: {os.path.basename(file_path)}", style='ManuscriptBody')

            except Exception as e:
                doc.add_paragraph(f"Error processing {os.path.basename(file_path)}: {str(e)}", style='ManuscriptBody')

            if i < len(components) - 1:  # Don't add page break after last component
                doc.add_page_break()
            section_num += 1

    # Create output filename
    output_name = manuscript_name.replace('.md', '.docx')
    output_path = f"final_manuscripts/{output_name}"

    # Ensure output directory exists
    os.makedirs('final_manuscripts', exist_ok=True)

    # Save document
    try:
        doc.save(output_path)
        print(f"✅ Successfully created: {output_path}")
        print(f"   Includes {len(components)} supporting components")
        return True
    except Exception as e:
        print(f"❌ Error saving {output_path}: {e}")
        return False

def main():
    """Main function to process all completed manuscripts"""

    print("🚀 Starting DOCX manuscript conversion...")
    print("=" * 60)

    # Process Fibromyalgia manuscript
    print("\n🔄 Processing Fibromyalgia Microbiome manuscript")
    print("-" * 50)
    fibro_main = 'Fibromyalgia_Microbiome_MetaAnalysis/final_manuscript.md'
    fibro_components = [
        ('validation_methodology', 'Fibromyalgia_Microbiome_MetaAnalysis/validation_methodology.md'),
        ('supplementary_materials', 'Fibromyalgia_Microbiome_MetaAnalysis/supplementary_materials.md'),
        ('PRISMA_flowchart', 'Fibromyalgia_Microbiome_MetaAnalysis/PRISMA_flowchart.md'),
        ('project_summary', 'Fibromyalgia_Microbiome_MetaAnalysis/project_summary.md')
    ]
    process_manuscript_directly('fibromyalgia_microbiome_manuscript.docx', fibro_main, fibro_components)

    # Process PPG manuscript
    print("\n🔄 Processing PPG HR Accuracy manuscript")
    print("-" * 50)
    ppg_main = 'ppg_hr_accuracy_meta_analysis/manuscript_draft.md'
    ppg_components = [
        ('protocol', 'ppg_hr_accuracy_meta_analysis/protocol.md'),
        ('search_strategy', 'ppg_hr_accuracy_meta_analysis/detailed_search_strategy.md'),
        ('data_extraction', 'ppg_hr_accuracy_meta_analysis/data_extraction_form.md'),
        ('forest_plot', 'ppg_hr_accuracy_meta_analysis/results/forest_plot_visualization.txt'),
        ('bland_altman', 'ppg_hr_accuracy_meta_analysis/results/bland_altman_plot.txt'),
        ('performance_table', 'ppg_hr_accuracy_meta_analysis/results/performance_comparison_table.md'),
        ('validation_report', 'ppg_hr_accuracy_meta_analysis/results/validation_report.md'),
        ('project_summary', 'ppg_hr_accuracy_meta_analysis/project_summary.md')
    ]
    process_manuscript_directly('ppg_hr_accuracy_manuscript.docx', ppg_main, ppg_components)

    print("\n" + "=" * 60)
    print("🎉 CONVERSION COMPLETE!")
    print("📁 Documents saved in: final_manuscripts/")

def process_manuscript_directly(output_filename, main_file_path, components):
    """Process manuscript directly with known components to avoid path issues"""

    if not os.path.exists(main_file_path):
        print(f"❌ Main manuscript not found: {main_file_path}")
        return False

    print(f"✅ Found main manuscript: {os.path.basename(main_file_path)}")
    print(f"📋 Will include {len(components)} supporting components")

    # Create document
    doc = Document()
    setup_document_styles(doc)

    # Title page
    title = doc.add_paragraph("Meta-Analysis Manuscript", style='ManuscriptTitle')
    subtitle = doc.add_paragraph("Complete Systematic Review and Meta-Analysis", style='ManuscriptSection')
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()

    # Main manuscript
    doc.add_paragraph("Main Manuscript", style='ManuscriptSection')

    try:
        with open(main_file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        convert_markdown_to_docx(content, doc)
    except Exception as e:
        print(f"❌ Error reading main manuscript: {e}")
        return False

    # Supporting components
    doc.add_page_break()
    doc.add_paragraph("Supporting Documentation", style='ManuscriptSection')

    for category, file_path in components:
        if os.path.exists(file_path):
            doc.add_paragraph(f".{category.replace('_', ' ').title()}", style='ManuscriptSection')

            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()

                if file_path.endswith('.txt'):
                    # Plain text
                    for line in content.split('\n'):
                        p = doc.add_paragraph(line, style='ManuscriptBody')
                else:
                    # Markdown
                    convert_markdown_to_docx(content, doc)

            except Exception as e:
                doc.add_paragraph(f"Error: {str(e)}", style='ManuscriptBody')
        else:
            doc.add_paragraph(f"File not found: {file_path}", style='ManuscriptBody')

    # Save with shorter filename
    output_path = f"final_manuscripts/{output_filename}"
    os.makedirs('final_manuscripts', exist_ok=True)

    try:
        doc.save(output_path)
        print(f"✅ SUCCESS: {output_path}")
        return True
    except Exception as e:
        print(f"❌ Failed to save: {e}")
        return False

if __name__ == '__main__':
    main()
