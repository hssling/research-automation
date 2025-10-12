#!/usr/bin/env python3
"""
PPG Heart Rate Accuracy Manuscript DOCX Converter
Converts Markdown manuscript to professional Word document format
"""

import re
from typing import List, Tuple
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

class ManuscriptConverter:
    def __init__(self, input_file: str, output_file: str):
        self.input_file = input_file
        self.output_file = output_file
        self.doc = Document()
        self.setup_styles()

    def setup_styles(self):
        """Setup custom styles for professional formatting"""
        styles = self.doc.styles

        # Title style
        title_style = styles.add_style('TitleStyle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.size = Pt(18)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_style.paragraph_format.space_after = Inches(0.5)

        # Authors style
        authors_style = styles.add_style('AuthorsStyle', WD_STYLE_TYPE.PARAGRAPH)
        authors_style.font.size = Pt(11)
        authors_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Section header
        section_style = styles.add_style('SectionStyle', WD_STYLE_TYPE.PARAGRAPH)
        section_style.font.size = Pt(14)
        section_style.font.bold = True
        section_style.paragraph_format.space_before = Inches(0.25)
        section_style.paragraph_format.space_after = Inches(0.15)

        # Subsection header
        subsection_style = styles.add_style('SubsectionStyle', WD_STYLE_TYPE.PARAGRAPH)
        subsection_style.font.size = Pt(12)
        subsection_style.font.bold = True
        subsection_style.paragraph_format.space_before = Inches(0.2)

        # Abstract style
        abstract_style = styles.add_style('AbstractStyle', WD_STYLE_TYPE.PARAGRAPH)
        abstract_style.font.size = Pt(10)
        abstract_style.font.italic = True
        abstract_style.paragraph_format.first_line_indent = Inches(0.25)

        # Body style
        body_style = styles.add_style('BodyStyle', WD_STYLE_TYPE.PARAGRAPH)
        body_style.font.size = Pt(11)
        body_style.font.name = 'Times New Roman'
        body_style.paragraph_format.line_spacing = 1.5
        body_style.paragraph_format.space_after = Pt(8)

        # Table title style
        table_title_style = styles.add_style('TableTitleStyle', WD_STYLE_TYPE.PARAGRAPH)
        table_title_style.font.size = Pt(11)
        table_title_style.font.bold = True
        table_title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Figure caption style
        figure_caption_style = styles.add_style('FigureCaptionStyle', WD_STYLE_TYPE.PARAGRAPH)
        figure_caption_style.font.size = Pt(11)
        figure_caption_style.font.italic = True
        figure_caption_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Reference style
        reference_style = styles.add_style('ReferenceStyle', WD_STYLE_TYPE.PARAGRAPH)
        reference_style.font.size = Pt(10)
        reference_style.paragraph_format.first_line_indent = Inches(0.3)

    def clean_markdown(self, text: str) -> str:
        """Clean common Markdown artifacts"""
        # Remove ** markup but preserve text
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        # Remove standalone asterisks
        text = text.replace('*', '')
        return text

    def parse_table(self, lines: List[str], start_index: int) -> Tuple[List[List[str]], int]:
        """Parse markdown table"""
        table_rows = []
        i = start_index

        while i < len(lines):
            line = lines[i].strip()
            if not line or not '|' in line:
                break
            # Skip separator row (contains only --- and |)
            if re.match(r'^[\s\|:\-]*$', line.replace(' ', '')):
                i += 1
                continue

            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells and any(cell for cell in cells):
                table_rows.append(cells)
            i += 1

        return table_rows, i - 1

    def add_table_to_doc(self, table_data: List[List[str]], caption: str = None):
        """Add table to document"""
        if not table_data:
            return

        table = self.doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        table.style = 'Table Grid'

        # Add data
        for row_idx, row_data in enumerate(table_data):
            for col_idx, cell_data in enumerate(row_data):
                cell_data = self.clean_markdown(cell_data)
                table.rows[row_idx].cells[col_idx].text = cell_data

                # Make header row bold
                if row_idx == 0:
                    for paragraph in table.rows[row_idx].cells[col_idx].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

        # Add caption
        if caption:
            self.doc.add_paragraph(caption, style='TableTitleStyle')

    def process_manuscript(self):
        """Process manuscript file"""
        with open(self.input_file, 'r', encoding='utf-8') as f:
            content = f.read()

        lines = content.split('\n')
        in_abstract = False
        skip_section = False

        i = 0
        while i < len(lines):
            line = lines[i].strip()

            if not line:
                if not skip_section:
                    self.doc.add_paragraph('')
                i += 1
                continue

            # Skip markdown separators
            if line == '---':
                i += 1
                continue

            # Main title
            if line.startswith('# '):
                title = line[2:]
                p = self.doc.add_paragraph(title, style='TitleStyle')

                # Process author information from next lines
                author_lines = []
                j = i + 1
                while j < len(lines) and (lines[j].strip() == '---' or
                                         '**Authors:**' in lines[j] or
                                         '¹Regulatory Research' in lines[j] or
                                         '**Corresponding Author:**' in lines[j] or
                                         'Email: framework@researchintegrity.org' in lines[j]):
                    author_line = lines[j].strip()
                    if author_line and author_line != '---':
                        author_lines.append(self.clean_markdown(author_line))
                    j += 1

                for author_line in author_lines:
                    self.doc.add_paragraph(author_line, style='AuthorsStyle')

                i = j - 1

            # Section headers
            elif line.startswith('## '):
                section_title = line[3:]

                # Skip certain sections
                if 'Plain Language Summary' in section_title:
                    skip_section = True
                    i += 1
                    continue
                else:
                    skip_section = False

                self.doc.add_paragraph(section_title, style='SectionStyle')
                if 'Abstract' in section_title:
                    in_abstract = True

            # Subsection headers
            elif line.startswith('### '):
                subsection_title = line[4:]
                self.doc.add_paragraph(subsection_title, style='SubsectionStyle')

            # Sub-subsection headers
            elif line.startswith('#### ') or line.startswith('##### '):
                subsubsection_title = line.lstrip('# ').strip()
                p = self.doc.add_paragraph(subsubsection_title, style='BodyStyle')
                for run in p.runs:
                    run.bold = True

            # Tables
            elif line.startswith('|') and '|' in line and not skip_section:
                table_data, end_index = self.parse_table(lines, i)
                if table_data:
                    # Get caption if exists
                    caption = None
                    if end_index + 1 < len(lines):
                        next_line = lines[end_index + 1].strip()
                        if '**Table' in next_line:
                            caption = self.clean_markdown(next_line)

                    self.add_table_to_doc(table_data, caption)
                    i = end_index + 1
                    if caption:
                        i += 1  # Skip caption line
                else:
                    i += 1
                continue

            # Figure placeholders
            elif '[' in line and 'Figure' in line and not skip_section:
                fig_caption = f"[Figure placeholder: {line}]"
                self.doc.add_paragraph(fig_caption, style='FigureCaptionStyle')

            # References section
            elif '**References**' in line or '## References' in line:
                self.doc.add_paragraph('References', style='SectionStyle')

            # Numbered lists or bullets
            elif line.startswith('- ') or re.match(r'^\d+\.', line):
                clean_line = self.clean_markdown(line)
                self.doc.add_paragraph(clean_line, style='BodyStyle')

            # Regular paragraphs
            elif not skip_section:
                clean_line = self.clean_markdown(line)
                if clean_line:
                    style = 'AbstractStyle' if in_abstract else 'BodyStyle'
                    self.doc.add_paragraph(clean_line, style=style)

            i += 1

    def save_document(self):
        """Save the document"""
        self.doc.save(self.output_file)
        print(f"✓ Manuscript converted successfully!")
        print(f"✓ Output: {self.output_file}")

def main():
    input_file = 'manuscript_draft.md'
    output_file = 'ppg_hr_accuracy_final_manuscript.docx'

    converter = ManuscriptConverter(input_file, output_file)
    print("🔄 Converting PPG Heart Rate Accuracy manuscript to DOCX...")
    converter.process_manuscript()
    converter.save_document()

if __name__ == '__main__':
    main()
