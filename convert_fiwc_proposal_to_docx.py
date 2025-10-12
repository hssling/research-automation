#!/usr/bin/env python3
"""
FIWC Proposal to DOCX Converter
Converts all ICMR FIWC Grant Proposal components into professional DOCX documents
"""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

def setup_document_styles(doc):
    """Setup document styles for professional grant proposal formatting"""
    # Title style
    title_style = doc.styles.add_style('ProposalTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Section header style
    section_style = doc.styles.add_style('ProposalSection', WD_STYLE_TYPE.PARAGRAPH)
    section_style.font.size = Pt(14)
    section_style.font.bold = True
    section_style.paragraph_format.space_before = Pt(12)

    # Subsection style
    subsection_style = doc.styles.add_style('ProposalSubsection', WD_STYLE_TYPE.PARAGRAPH)
    subsection_style.font.size = Pt(12)
    subsection_style.font.bold = True
    subsection_style.paragraph_format.space_before = Pt(8)

    # Normal text style
    body_style = doc.styles.add_style('ProposalBody', WD_STYLE_TYPE.PARAGRAPH)
    body_style.font.size = Pt(11)
    body_style.paragraph_format.line_spacing = 1.15

    # Table/text style
    table_style = doc.styles.add_style('ProposalTable', WD_STYLE_TYPE.PARAGRAPH)
    table_style.font.size = Pt(10)

def convert_markdown_to_docx(content, doc):
    """Convert markdown content to professional DOCX format"""
    lines = content.split('\n')
    current_list = None
    in_code_block = False
    code_content = []
    in_table = False
    table_rows = []

    for line in lines:
        line = line.rstrip()

        if not line.strip() and not in_code_block and not in_table:
            if current_list:
                current_list = None
            else:
                doc.add_paragraph('')
            continue

        # Code block handling
        if line.strip().startswith('```'):
            if in_code_block:
                # End of code block
                code_text = '\n'.join(code_content)
                p = doc.add_paragraph(code_text)
                p.style = doc.styles['ProposalBody']
                p.paragraph_format.left_indent = Inches(0.25)
                in_code_block = False
                code_content = []
            else:
                # Start of code block
                in_code_block = True
            continue

        if in_code_block:
            code_content.append(line)
            continue

        # Headers
        if line.startswith('# '):
            p = doc.add_paragraph(line[2:], style='ProposalTitle')
        elif line.startswith('## '):
            p = doc.add_paragraph(line[3:], style='ProposalSection')
        elif line.startswith('### '):
            p = doc.add_paragraph(line[4:], style='ProposalSubsection')
        elif line.startswith('#### ') or line.startswith('##### ') or line.startswith('###### '):
            p = doc.add_paragraph(line.lstrip('# '), style='ProposalBody')
            p.runs[0].bold = True

        # Table handling
        elif '|' in line and not in_table:
            # Check if this looks like a table
            parts = [p.strip() for p in line.split('|') if p.strip()]
            if len(parts) > 1 and not line.startswith('-'):
                in_table = True
                table_rows.append(parts)

        elif in_table and '|' in line:
            if line.strip().startswith('|---') or line.strip().startswith('|---'):
                # Separator line, skip
                continue
            else:
                parts = [p.strip() for p in line.split('|') if p.strip()]
                table_rows.append(parts)

        elif in_table and not '|' in line and line.strip():
            # End of table, create the table
            if table_rows:
                table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                table.style = 'Table Grid'

                for i, row in enumerate(table_rows):
                    for j, cell_text in enumerate(row):
                        if i < len(table.rows) and j < len(table.rows[i].cells):
                            cell = table.rows[i].cells[j]
                            cell.paragraphs[0].text = cell_text
                            cell.paragraphs[0].style = 'ProposalTable'

            table_rows = []
            in_table = False
            # Process current line as regular text
            if line.strip():
                p = doc.add_paragraph(line, style='ProposalBody')

        # Lists
        elif not in_table and (line.strip().startswith('- ') or line.strip().startswith('* ')):
            if current_list is None:
                current_list = doc.add_paragraph(style='ProposalBody')
            else:
                current_list = doc.add_paragraph(style='ProposalBody')
            current_list.add_run(line.strip()[2:])
            current_list.paragraph_format.left_indent = Inches(0.25)

        elif not in_table and line.strip() and (line.strip()[0].isdigit() and '. ' in line):
            if current_list is None:
                current_list = doc.add_paragraph(style='ProposalBody')
            else:
                current_list = doc.add_paragraph(style='ProposalBody')
            current_list.add_run(line.strip())
            current_list.paragraph_format.left_indent = Inches(0.25)

        else:
            if not in_table:
                # Regular paragraph
                p = doc.add_paragraph(line, style='ProposalBody')
                current_list = None

def create_fiwc_document(file_path, title, output_path):
    """Create a professional DOCX document from a FIWC proposal markdown file"""

    if not os.path.exists(file_path):
        print(f"❌ File not found: {file_path}")
        return False

    print(f"📄 Converting: {os.path.basename(file_path)}")

    # Create document
    doc = Document()
    setup_document_styles(doc)

    # Add title page
    title_para = doc.add_paragraph(title, style='ProposalTitle')
    title_para.paragraph_format.space_after = Pt(24)

    # Add header info for official documents
    header_info = doc.add_paragraph("ICMR First-in-World Challenge (FIWC) Grant Proposal", style='ProposalBody')
    header_info.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_info.paragraph_format.space_after = Pt(12)

    doc.add_paragraph("Principal Investigator: Dr. [Name] | Institution: [Institution]", style='ProposalBody')
    doc.add_paragraph("Date: [Current Date] | Grant Category: FIWC - First-in-World Innovation", style='ProposalBody')

    doc.add_page_break()

    # Convert content
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # Add table of contents placeholder
        toc = doc.add_paragraph("Table of Contents", style='ProposalSection')
        toc.paragraph_format.space_after = Pt(12)
        doc.add_paragraph("[TOC will be generated in final document]", style='ProposalBody')
        doc.add_page_break()

        convert_markdown_to_docx(content, doc)

    except Exception as e:
        print(f"❌ Error processing {file_path}: {e}")
        doc.add_paragraph(f"Error loading document: {str(e)}", style='ProposalBody')
        return False

    # Save document
    try:
        doc.save(output_path)
        print(f"✅ Created: {output_path}")
        return True
    except Exception as e:
        print(f"❌ Error saving {output_path}: {e}")
        return False

def main():
    """Convert all FIWC proposal files to DOCX format"""

    print("🚀 Converting ICMR FIWC Proposal to DOCX Format...")
    print("=" * 60)

    # Ensure output directory exists
    os.makedirs('ICMR_FIWC_Proposal_DOCX', exist_ok=True)

    # Define files to convert
    proposal_files = {
        # Main proposal components
        '01_main_proposal/FIWC_Proposal_Autonomous_Research_System.md':
            'FIWC_Proposal_Main_Document.docx',

        '02_technical_methodology/technical_methodology.md':
            'FIWC_Proposal_Technical_Methodology.docx',

        '03_budget_timeline/project_timeline_gantt.md':
            'FIWC_Proposal_Project_Timeline.docx',

        '03_budget_timeline/budget_justification.md':
            'FIWC_Proposal_Budget_Justification.docx',

        '05_team_credentials/principal_investigator_resume.md':
            'FIWC_Proposal_Principal_Investigator_CV.docx',

        '06_regulatory_compliance/regulatory_compliance_framework.md':
            'FIWC_Proposal_Regulatory_Compliance.docx',

        '07_appendices/compliance_validation_checklist.md':
            'FIWC_Proposal_Validation_Checklist.docx',

        # Root documents
        'README.md':
            'FIWC_Proposal_Project_Overview.docx'
    }

    success_count = 0
    total_files = len(proposal_files)

    for input_file, output_file in proposal_files.items():
        input_path = f"ICMR_FIWC_Proposal/{input_file}"
        output_path = f"ICMR_FIWC_Proposal_DOCX/{output_file}"

        # Create meaningful title from filename
        title = output_file.replace('FIWC_Proposal_', '').replace('.docx', '').replace('_', ' ').title()
        if 'Main Document' in title:
            title = "Autonomous Research Automation System - FIWC Grant Proposal"
        elif 'Technical Methodology' in title:
            title = "Technical Methodology and System Architecture"
        elif 'Project Timeline' in title:
            title = "Project Timeline and Implementation Plan"
        elif 'Budget Justification' in title:
            title = "Detailed Budget Justification"
        elif 'Principal Investigator Cv' in title:
            title = "Principal Investigator Curriculum Vitae"
        elif 'Regulatory Compliance' in title:
            title = "Regulatory Compliance Framework"
        elif 'Validation Checklist' in title:
            title = "Proposal Compliance Validation"
        elif 'Project Overview' in title:
            title = "ICMR FIWC Proposal Project Overview"

        if create_fiwc_document(input_path, title, output_path):
            success_count += 1

    print("\n" + "=" * 60)
    print(f"🎉 CONVERSION COMPLETE!")
    print(f"✅ Successfully converted: {success_count}/{total_files} files")
    print(f"📁 DOCX documents saved in: ICMR_FIWC_Proposal_DOCX/")

    if success_count == total_files:
        print("\n📋 All FIWC proposal components converted to professional MS Word DOCX format")
        print("📄 Ready for ICMR submission with proper academic formatting")
    else:
        print(f"\n⚠️  {total_files - success_count} files had conversion issues - please check output")

if __name__ == '__main__':
    main()
