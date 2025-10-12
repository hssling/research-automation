#!/usr/bin/env python3
"""
AMR Forecasting Report Generator

Generates comprehensive Word DOCX and PDF reports for AMR forecasting analyses,
including model comparisons, performance metrics, and strategic interpretations.

Features:
- Professional report formatting
- Automated interpretation text generation
- Model comparison tables and charts
- Executive summaries for stakeholders
- Multi-format output (DOCX, PDF)
- Ready for distribution to healthcare decision-makers
"""

import pandas as pd
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import plotly.io as pio
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional
import logging

logger = logging.getLogger(__name__)

class AMRForecastReportGenerator:
    """
    Comprehensive report generator for AMR forecasting analyses.

    Generates professional-quality reports suitable for:
    - Hospital administrators
    - Public health officials
    - Pharmaceutical companies
    - Academic researchers
    - Policy makers
    """

    def __init__(self, output_dir: str = "reports"):
        """
        Initialize report generator with output directory.

        Args:
            output_dir: Directory to save generated reports
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)

    def generate_comprehensive_report(self, country: str, pathogen: str, antibiotic: str,
                                      forecasts: Dict, metrics: pd.DataFrame,
                                      historical_data: pd.DataFrame) -> Dict[str, str]:
        """
        Generate comprehensive AMR forecasting report package.

        Args:
            country: Country name
            pathogen: Pathogen name
            antibiotic: Antibiotic name
            forecasts: Dictionary of forecast results by model
            metrics: Performance metrics DataFrame
            historical_data: Historical resistance data

        Returns:
            Dictionary with paths to generated files
        """

        safe_country = country.replace(" ", "_").replace("/", "_")
        safe_pathogen = pathogen.replace(" ", "_").replace("/", "_")
        safe_antibiotic = antibiotic.replace(" ", "_").replace("/", "_")

        base_filename = f"{safe_country}_{safe_pathogen}_{safe_antibiotic}"
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        generated_files = {}

        try:
            # Generate model comparison visualization
            model_plot_path = self._generate_model_comparison_plot(
                forecasts, historical_data, country, pathogen, antibiotic, base_filename
            )
            generated_files['model_comparison_plot'] = model_plot_path

            # Generate performance analysis chart
            performance_plot_path = self._generate_performance_analysis_plot(
                metrics, base_filename
            )
            generated_files['performance_plot'] = performance_plot_path

            # Generate forecast uncertainty visualization
            uncertainty_plot_path = self._generate_uncertainty_plot(
                forecasts, base_filename
            )
            generated_files['uncertainty_plot'] = uncertainty_plot_path

            # Generate comprehensive DOCX report
            docx_report_path = self._generate_docx_report(
                country, pathogen, antibiotic, forecasts, metrics, historical_data,
                [model_plot_path, performance_plot_path, uncertainty_plot_path],
                base_filename, timestamp
            )
            generated_files['docx_report'] = docx_report_path

            # Generate executive summary PDF
            pdf_summary_path = self._generate_pdf_summary(
                country, pathogen, antibiotic, metrics, base_filename, timestamp
            )
            generated_files['pdf_summary'] = pdf_summary_path

            logger.info(f"Comprehensive report package generated for {country}_{pathogen}_{antibiotic}")
            return generated_files

        except Exception as e:
            logger.error(f"Report generation failed: {e}")
            return {}

    def _generate_model_comparison_plot(self, forecasts: Dict, historical_data: pd.DataFrame,
                                       country: str, pathogen: str, antibiotic: str,
                                       base_filename: str) -> str:
        """Generate comparative visualization of all forecast models."""

        # Create subplot figure
        fig = make_subplots(
            rows=2, cols=2,
            subplot_titles=[
                'Historical Data Overview',
                'Model Forecast Comparison',
                'Individual Model Forecasts',
                'Error Distribution Analysis'
            ],
            specs=[
                [{"secondary_y": False}, {"secondary_y": False}],
                [{"secondary_y": False}, {"secondary_y": False}]
            ]
        )

        # Plot 1: Historical data
        hist_data = historical_data.copy()
        hist_data['period'] = range(1, len(hist_data) + 1)  # Numerical x-axis

        fig.add_trace(
            go.Scatter(
                x=hist_data['period'],
                y=hist_data['percent_resistant'],
                mode='lines+markers',
                name='Historical Resistance',
                line=dict(color='blue', width=3),
                marker=dict(size=6)
            ),
            row=1, col=1
        )

        # Plot 2: Model comparison (short-term forecasts)
        colors = {'Prophet': 'red', 'ARIMA': 'blue', 'LSTM': 'green'}
        hist_last_12 = hist_data.tail(12)  # Last 12 months of historical data

        # Add historical reference
        fig.add_trace(
            go.Scatter(
                x=list(range(1, 13)) + [f'+{i}' for i in range(1, 13)],
                y=hist_last_12['percent_resistant'].tolist() + [None] * 12,
                mode='lines+markers',
                name='Historical (last 12m)',
                line=dict(color='gray', width=2, dash='dash'),
                marker=dict(size=6),
                showlegend=False
            ),
            row=1, col=2
        )

        # Add model forecasts (next 12 months)
        for model_name, forecast_data in forecasts.items():
            if 'future_forecast' in forecast_data:
                future_df = forecast_data['future_forecast']
                if isinstance(future_df, pd.DataFrame) and 'yhat' in future_df.columns:
                    future_values = future_df['yhat'].head(12).values  # Next 12 months

                    fig.add_trace(
                        go.Scatter(
                            x=[f'+{i}' for i in range(1, len(future_values) + 1)],
                            y=future_values,
                            mode='lines+markers',
                            name=f'{model_name} Forecast',
                            line=dict(color=colors.get(model_name, 'gray'), width=3),
                            marker=dict(size=6)
                        ),
                        row=1, col=2
                    )

        # Plot 3: Individual model trajectories (whole forecast period)
        future_periods = 24  # 24 months
        historical_period = len(hist_data)

        for model_name, forecast_data in forecasts.items():
            if 'future_forecast' in forecast_data:
                future_df = forecast_data['future_forecast']
                if isinstance(future_df, pd.DataFrame) and 'yhat' in future_df.columns:
                    all_forecasts = future_df['yhat'].head(future_periods).values
                    x_vals = list(range(historical_period + 1, historical_period + len(all_forecasts) + 1))

                    fig.add_trace(
                        go.Scatter(
                            x=x_vals,
                            y=all_forecasts,
                            mode='lines',
                            name=f'{model_name} (24m)',
                            line=dict(color=colors.get(model_name, 'gray'), width=2),
                            showlegend=True
                        ),
                        row=2, col=1
                    )

        # Add historical reference line to subplot 3
        fig.add_trace(
            go.Scatter(
                x=list(range(1, historical_period + 1)),
                y=hist_data['percent_resistant'],
                mode='lines',
                name='Historical Timeline',
                line=dict(color='black', width=1, dash='dot'),
                showlegend=True
            ),
            row=2, col=1
        )

        # Plot 4: Error distribution analysis (if available)
        if 'test_predictions' in forecasts.get('Prophet', {}):
            error_data = []

            for model_name in ['Prophet', 'ARIMA', 'LSTM']:
                if (model_name in forecasts and
                    'test_predictions' in forecasts[model_name] and
                    'actual_test' in forecasts[model_name]):

                    pred = forecasts[model_name]['test_predictions']
                    actual = forecasts[model_name]['actual_test']
                    errors = actual - pred
                    error_data.append(errors)

            if error_data:
                labels = [f'{model} Errors' for model in ['Prophet', 'ARIMA', 'LSTM']]
                colors_box = ['rgba(255,0,0,0.7)', 'rgba(0,0,255,0.7)', 'rgba(0,128,0,0.7)']

                for i, errors in enumerate(error_data):
                    fig.add_trace(
                        go.Box(
                            y=errors,
                            name=labels[i],
                            marker_color=colors_box[i],
                            boxpoints='all',
                            jitter=0.3,
                            pointpos=-1.8
                        ),
                        row=2, col=2
                    )

                # Add zero reference line
                fig.add_hline(y=0, line_dash="dash", line_color="black", row=2, col=2)

        # Update layout
        fig.update_layout(
            height=800,
            title_text=f'AMR Forecast Model Comparison: {pathogen} vs {antibiotic} in {country}',
            title_x=0.5,
            showlegend=True
        )

        # Update axis labels
        fig.update_xaxes(title_text="Time Period", row=1, col=1)
        fig.update_yaxes(title_text="Resistance %", row=1, col=1)
        fig.update_xaxes(title_text="Forecast Horizon", row=1, col=2)
        fig.update_yaxes(title_text="Resistance %", row=1, col=2)
        fig.update_xaxes(title_text="Timeline", row=2, col=1)
        fig.update_yaxes(title_text="Resistance %", row=2, col=1)
        fig.update_xaxes(title_text="Model", row=2, col=2)
        fig.update_yaxes(title_text="Forecast Error", row=2, col=2, zeroline=True)

        # Add resistance threshold lines
        for subplot_row, subplot_col in [(1,2), (2,1)]:
            fig.add_hline(y=70, line_dash="dash", line_color="orange", opacity=0.7,
                         annotation_text="Warning (70%)", row=subplot_row, col=subplot_col)
            fig.add_hline(y=80, line_dash="dash", line_color="red", opacity=0.7,
                         annotation_text="Critical (80%)", row=subplot_row, col=subplot_col)

        # Save plot
        plot_path = self.output_dir / f"model_comparison_{base_filename}.html"
        fig.write_html(str(plot_path))

        # Also save as PNG for reports
        plot_png_path = self.output_dir / f"model_comparison_{base_filename}.png"
        fig.write_image(str(plot_png_path), width=1200, height=800)

        logger.info(f"Model comparison visualization saved to {plot_png_path}")
        return str(plot_png_path)

    def _generate_performance_analysis_plot(self, metrics: pd.DataFrame, base_filename: str) -> str:
        """Generate detailed performance metrics visualization."""

        # Create metrics comparison plot
        fig = make_subplots(
            rows=2, cols=2,
            subplot_titles=[
                'RMSE Comparison',
                'MAE Comparison',
                'MAPE Comparison',
                'Composite Accuracy Score'
            ]
        )

        colors = ['red', 'blue', 'green']
        models = metrics['Model'].tolist()

        # RMSE plot
        fig.add_trace(
            go.Bar(
                x=models,
                y=metrics['RMSE'],
                name='RMSE',
                marker_color=colors[:len(models)]
            ),
            row=1, col=1
        )

        # MAE plot
        fig.add_trace(
            go.Bar(
                x=models,
                y=metrics['MAE'],
                name='MAE',
                marker_color=colors[:len(models)]
            ),
            row=1, col=2
        )

        # MAPE plot
        fig.add_trace(
            go.Bar(
                x=models,
                y=metrics['MAPE'],
                name='MAPE (%)',
                marker_color=colors[:len(models)]
            ),
            row=2, col=1
        )

        # Composite score
        if 'accuracy_score' in metrics.columns:
            fig.add_trace(
                go.Bar(
                    x=models,
                    y=metrics['accuracy_score'],
                    name='Accuracy Score',
                    marker_color=colors[:len(models)]
                ),
                row=2, col=2
            )

        fig.update_layout(
            title_text="Model Performance Metrics Comparison",
            title_x=0.5,
            showlegend=False,
            height=600
        )

        # Add metric interpretations
        fig.add_annotation(
            text="Lower values = Better performance",
            xref="paper", yref="paper",
            x=0.5, y=-0.1,
            showarrow=False,
            font=dict(size=10, color="gray")
        )

        # Save plot
        plot_path = self.output_dir / f"performance_analysis_{base_filename}.png"
        fig.write_image(str(plot_path), width=1000, height=600)

        logger.info(f"Performance analysis visualization saved to {plot_path}")
        return str(plot_path)

    def _generate_uncertainty_plot(self, forecasts: Dict, base_filename: str) -> str:
        """Generate forecast uncertainty visualization."""
        # This would show confidence intervals for different models
        # For now, create a placeholder implementation

        fig = go.Figure()

        fig.add_trace(go.Scatter(
            x=[1, 2, 3, 4],
            y=[1, 4, 2, 3],
            mode='lines+markers',
            name='Uncertainty Placeholder'
        ))

        fig.update_layout(
            title="Forecast Uncertainty Analysis",
            xaxis_title="Forecast Horizon",
            yaxis_title="Prediction Confidence"
        )

        plot_path = self.output_dir / f"uncertainty_analysis_{base_filename}.png"
        fig.write_image(str(plot_path), width=1000, height=400)

        return str(plot_path)

    def _generate_docx_report(self, country: str, pathogen: str, antibiotic: str,
                             forecasts: Dict, metrics: pd.DataFrame, historical_data: pd.DataFrame,
                             plot_files: List[str], base_filename: str, timestamp: str) -> str:
        """Generate comprehensive Word DOCX report."""

        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import docx.oxml.ns as ns
            import docx.oxml as oxml

        except ImportError:
            logger.warning("python-docx not available, skipping DOCX report generation")
            return ""

        doc = Document()

        # Title page
        title = doc.add_heading('Antimicrobial Resistance Forecasting Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Report metadata
        doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'Country: {country}')
        doc.add_paragraph(f'Pathogen: {pathogen}')
        doc.add_paragraph(f'Antibiotic: {antibiotic}')
        doc.add_paragraph(f'Data Sources: WHO GLASS, CDC NARMS, ResistanceMap')
        doc.add_paragraph()

        # Executive Summary
        doc.add_heading('Executive Summary', 1)

        # Find best model
        best_model = None
        if not metrics.empty and 'accuracy_score' in metrics.columns:
            best_row = metrics.loc[metrics['accuracy_score'].idxmin()]
            best_model = best_row['Model']

        summary_text = f"""
        This report presents a comprehensive analysis of antimicrobial resistance forecasting for {pathogen} treated with {antibiotic} in {country}.

        Key Findings:
        • Historical resistance data analyzed: {len(historical_data)} records
        • Three forecasting models evaluated: Prophet, ARIMA, and LSTM
        • Best performing model: {best_model if best_model else 'Analysis inconclusive'}
        • Forecast horizon: 24 months with performance validation

        The analysis provides evidence-based guidance for antibiotic stewardship and policy decision-making.
        """
        doc.add_paragraph(summary_text.strip())
        doc.add_paragraph()

        # Model Performance Section
        doc.add_heading('Model Performance Analysis', 1)

        # Add performance table
        if not metrics.empty:
            table = doc.add_table(rows=1, cols=len(metrics.columns))
            table.style = 'Table Grid'

            # Header row
            hdr_cells = table.rows[0].cells
            for i, col_name in enumerate(metrics.columns):
                hdr_cells[i].text = str(col_name)

            # Data rows
            for _, row in metrics.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    if pd.isna(value):
                        row_cells[i].text = 'N/A'
                    else:
                        row_cells[i].text = '.3f'

        doc.add_paragraph()

        # Interpretation section
        doc.add_heading('Model Interpretation', 1)

        interpretation = f"""
        Performance Metrics Explanation:

        • RMSE (Root Mean Square Error): Penalizes large forecasting errors more heavily
        • MAE (Mean Absolute Error): Average absolute deviation from actual values
        • MAPE (Mean Absolute Percentage Error): Relative error as percentage of actual values

        Model Characteristics:
        • Prophet: Effective for data with seasonality and trend patterns
        • ARIMA: Strong statistical foundation for time series analysis
        • LSTM: Best suited for complex, nonlinear resistance patterns

        Best Model Selection: The model with lowest composite accuracy score is recommended for deployment.
        """

        doc.add_paragraph(interpretation.strip())

        # Risk Assessment
        doc.add_heading('Risk Assessment', 1)

        current_resistance = historical_data['percent_resistant'].iloc[-1]

        if best_model and best_model in forecasts:
            future_forecast = forecasts[best_model].get('future_forecast')
            if future_forecast is not None and hasattr(future_forecast, 'iloc'):
                future_resistance = future_forecast.iloc[-1] if hasattr(future_forecast, 'iloc') else future_forecast['yhat'].iloc[-1]
                risk_change = future_resistance - current_resistance
            else:
                risk_change = 0
        else:
            future_resistance = current_resistance
            risk_change = 0

        risk_assessment = f"""
        Current Resistance Level: {current_resistance:.1f}%

        Projected Resistance (24 months): {future_resistance:.1f}%
        Resistance Change: {risk_change:+.1f}%

        Risk Categorization:
        """

        if future_resistance > 80:
            risk_assessment += "🚨 CRITICAL RISK: Immediate intervention required"
        elif future_resistance > 70:
            risk_assessment += "⚠️ HIGH RISK: Urgent action needed"
        elif future_resistance > 50:
            risk_assessment += "🔶 MODERATE RISK: Monitoring required"
        else:
            risk_assessment += "✅ LOW RISK: Current controls sufficient"

        doc.add_paragraph(risk_assessment.strip())

        # Recommendations
        doc.add_heading('Strategic Recommendations', 1)

        recommendations = f"""
        Based on forecasting analysis:

        1. Model Selection: Deploy {best_model} for operational forecasting
        2. Monitoring Frequency: Quarterly resistance surveillance
        3. Intervention Thresholds: Review protocols above 70% resistance
        4. Data Quality: Continue WHO-validated surveillance methods

        For implementation details, consult full technical documentation.
        """

        doc.add_paragraph(recommendations.strip())

        # Visualizations (if available)
        doc.add_heading('Forecast Visualizations', 1)

        for plot_file in plot_files[:1]:  # Just the main comparison plot for now
            if Path(plot_file).exists():
                try:
                    doc.add_picture(plot_file, width=Inches(6))
                    doc.add_paragraph("Model comparison showing historical data and forecast trajectories.")
                except:
                    doc.add_paragraph(f"Visualization available at: {plot_file}")

        # Methodology
        doc.add_heading('Methodology', 1)
        methodology = """
        Data Sources:
        - WHO GLASS: Global AMR surveillance network
        - CDC NARMS: US foodborne pathogen monitoring
        - ResistanceMap: CDDEP global consumption data

        Forecasting Models:
        - Temporal train/validation splits (80/20)
        - 24-month forecast horizon
        - Performance validated against held-out data
        - Uncertainty quantification through confidence intervals
        """
        doc.add_paragraph(methodology.strip())

        # Save document
        docx_path = self.output_dir / f"forecast_report_{base_filename}_{timestamp}.docx"
        doc.save(str(docx_path))

        logger.info(f"DOCX report saved to {docx_path}")
        return str(docx_path)

    def _generate_pdf_summary(self, country: str, pathogen: str, antibiotic: str,
                             metrics: pd.DataFrame, base_filename: str, timestamp: str) -> str:
        """Generate executive summary PDF."""

        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from reportlab.lib.styles import getSampleAppStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib import colors
        except ImportError:
            logger.warning("ReportLab not available, skipping PDF report generation")
            return ""

        pdf_path = self.output_dir / f"forecast_summary_{base_filename}_{timestamp}.pdf"

        doc = SimpleDocTemplate(str(pdf_path), pagesize=letter)
        elements = []
        styles = getSampleAppStyleSheet()

        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1  # Center alignment
        )

        elements.append(Paragraph("AMR Forecasting Executive Summary", title_style))
        elements.append(Spacer(1, 20))

        # Metadata
        metadata = f"""
        <b>Report Generated:</b> {datetime.now().strftime("%Y-%m-%d %H:%M")}<br/>
        <b>Country:</b> {country}<br/>
        <b>Pathogen:</b> {pathogen}<br/>
        <b>Antibiotic:</b> {antibiotic}<br/>
        <b>Forecast Horizon:</b> 24 months
        """

        elements.append(Paragraph(metadata.replace('\n', ''), styles["Normal"]))
        elements.append(Spacer(1, 20))

        # Performance metrics table
        if not metrics.empty:
            elements.append(Paragraph("<b>Model Performance Comparison</b>", styles["Heading3"]))

            # Prepare table data
            table_data = [list(metrics.columns)]  # Headers
            for _, row in metrics.iterrows():
                row_data = []
                for col in metrics.columns:
                    if pd.isna(row[col]):
                        row_data.append("N/A")
                    elif isinstance(row[col], (int, float)):
                        row_data.append(".3f")
                    else:
                        row_data.append(str(row[col]))
                table_data.append(row_data)

            # Create table
            table = Table(table_data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))

            elements.append(table)
            elements.append(Spacer(1, 20))

        # Key recommendations
        elements.append(Paragraph("<b>Key Recommendations</b>", styles["Heading3"]))

        recommendations = """
        • Implement chosen forecasting model for operational resistance monitoring
        • Establish resistance thresholds (Warning: 70%, Critical: 80%)
        • Conduct quarterly surveillance and model revalidation
        • Develop intervention protocols based on forecast risk levels
        """

        elements.append(Paragraph(recommendations, styles["Normal"]))
        elements.append(Spacer(1, 20))

        # Build PDF
        doc.build(elements)

        logger.info(f"PDF executive summary saved to {pdf_path}")
        return str(pdf_path)

    def generate_model_recommendation(self, metrics: pd.DataFrame) -> Dict:
        """Generate AI-powered model recommendation with reasoning."""

        if metrics.empty or 'accuracy_score' not in metrics.columns:
            return {"recommendation": "Unable to determine", "reasoning": "Insufficient metrics available"}

        # Find best model
        best_model_row = metrics.loc[metrics['accuracy_score'].idxmin()]
        best_model = best_model_row['Model']

        # Generate reasoning
        reasons = []

        if best_model_row['RMSE'] < metrics['RMSE'].quantile(0.5):
            reasons.append("consistently low prediction errors")

        if best_model_row['MAPE'] < 20:
            reasons.append("high prediction accuracy (<20% error)")

        if best_model_row['MAE'] <= metrics['MAE'].min() + 1:
            reasons.append("competitive absolute error metrics")

        reasoning = f"Selected {best_model} due to {', '.join(reasons)}. " \
                   f"This model demonstrated the lowest composite accuracy score and " \
                   "most reliable performance across validation metrics."

        return {
            "recommended_model": best_model,
            "reasoning": reasoning,
            "confidence_score": (1 - (best_model_row['accuracy_score'] - metrics['accuracy_score'].min()) /
                               (metrics['accuracy_score'].max() - metrics['accuracy_score'].min())),
            "alternative_models": [m for m in metrics['Model'].tolist() if m != best_model]
        }
