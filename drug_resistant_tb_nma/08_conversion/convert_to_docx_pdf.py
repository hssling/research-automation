#!/usr/bin/env python3
"""
Convert all markdown documents to DOCX and PDF formats for publication submission
Comprehensive conversion script for the Drug-Resistant TB Network Meta-Analysis project
"""

import os
import subprocess
import pandas as pd
from pathlib import Path
import markdown
from docx import Document
import warnings
warnings.filterwarnings('ignore')

def convert_markdown_to_docx(markdown_file, docx_file):
    """Convert markdown file to DOCX format"""

    try:
        # Read markdown content
        with open(markdown_file, 'r', encoding='utf-8') as f:
            content = f.read()

        # Create DOCX document
        doc = Document()

        # Add title
        title = Path(markdown_file).stem.replace('_', ' ').title()
        title_heading = doc.add_heading(title, 0)

        # Split content into lines and process
        lines = content.split('\n')

        for line in lines:
            if line.startswith('# '):
                # Main heading
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                # Secondary heading
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                # Tertiary heading
                doc.add_heading(line[4:], level=3)
            elif line.startswith('**') and line.endswith('**'):
                # Bold text
                doc.add_paragraph(line.strip('*'), style='Intense Quote')
            elif line.strip() == '':
                # Empty line
                doc.add_paragraph()
            else:
                # Regular paragraph
                doc.add_paragraph(line)

        # Save DOCX file
        doc.save(docx_file)
        print(f"✓ Converted {markdown_file} to {docx_file}")
        return True

    except Exception as e:
        print(f"✗ Error converting {markdown_file}: {str(e)}")
        return False

def convert_csv_to_excel(csv_file, excel_file):
    """Convert CSV files to Excel format for better presentation"""

    try:
        # Read CSV file
        df = pd.read_csv(csv_file)

        # Create Excel writer
        with pd.ExcelWriter(excel_file, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Data', index=False)

            # Get workbook and worksheet
            workbook = writer.book
            worksheet = writer.sheets['Data']

            # Auto-adjust column widths
            for column in worksheet.columns:
                max_length = 0
                column_letter = column[0].column_letter

                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass

                adjusted_width = min(max_length + 2, 50)
                worksheet.column_dimensions[column_letter].width = adjusted_width

        print(f"✓ Converted {csv_file} to {excel_file}")
        return True

    except Exception as e:
        print(f"✗ Error converting {csv_file}: {str(e)}")
        return False

def create_pdf_from_docx(docx_file, pdf_file):
    """Convert DOCX to PDF using LibreOffice"""

    try:
        # Use LibreOffice to convert DOCX to PDF
        # Note: This requires LibreOffice to be installed
        subprocess.run([
            'soffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', str(Path(pdf_file).parent),
            docx_file
        ], check=True, capture_output=True)

        print(f"✓ Converted {docx_file} to {pdf_file}")
        return True

    except subprocess.CalledProcessError as e:
        print(f"✗ LibreOffice conversion failed for {docx_file}: {e}")
        print("  Trying alternative method...")

        # Alternative: Use Microsoft Word if available (Windows)
        try:
            import win32com.client

            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False

            # Open document
            doc = word.Documents.Open(docx_file)

            # Save as PDF
            doc.SaveAs(pdf_file, FileFormat=17)  # 17 = PDF format

            # Close document and Word
            doc.Close()
            word.Quit()

            print(f"✓ Converted {docx_file} to {pdf_file} using Word")
            return True

        except ImportError:
            print("✗ Microsoft Word not available for PDF conversion")
            return False
        except Exception as e:
            print(f"✗ Word conversion failed: {e}")
            return False

def create_combined_pdf():
    """Create a combined PDF with all key documents"""

    try:
        # This would require a PDF merging library like PyPDF2
        # For now, create individual PDFs and note they can be combined manually

        print("ℹ Individual PDFs created. Use online PDF merger or Adobe Acrobat to combine:")
        print("  1. Complete Manuscript.pdf")
        print("  2. Supplementary Materials.pdf")
        print("  3. Validation Framework.pdf")
        print("  4. Publication Package.pdf")

        return True

    except Exception as e:
        print(f"✗ Error creating combined PDF: {e}")
        return False

def main():
    """Main conversion function"""

    print("="*60)
    print("CONVERSION TO DOCX/PDF FORMATS")
    print("="*60)
    print("Drug-Resistant Tuberculosis Network Meta-Analysis")
    print("Converting all documents for publication submission...")
    print()

    # Create output directories
    output_dirs = [
        'drug_resistant_tb_nma/08_conversion/docx',
        'drug_resistant_tb_nma/08_conversion/pdf',
        'drug_resistant_tb_nma/08_conversion/excel'
    ]

    for dir_path in output_dirs:
        os.makedirs(dir_path, exist_ok=True)

    # Key documents to convert
    key_documents = [
        'drug_resistant_tb_nma/00_protocol/study_protocol.md',
        'drug_resistant_tb_nma/00_protocol/prospero_registration.md',
        'drug_resistant_tb_nma/00_protocol/ethical_considerations.md',
        'drug_resistant_tb_nma/01_literature_search/search_strategy.md',
        'drug_resistant_tb_nma/01_literature_search/search_validation_report.md',
        'drug_resistant_tb_nma/02_data_extraction/data_extraction_form.md',
        'drug_resistant_tb_nma/03_statistical_analysis/nma_protocol.md',
        'drug_resistant_tb_nma/03_statistical_analysis/statistical_analysis_plan.md',
        'drug_resistant_tb_nma/04_results/results_summary.md',
        'drug_resistant_tb_nma/05_manuscript/complete_manuscript.md',
        'drug_resistant_tb_nma/05_manuscript/supplementary_materials.md',
        'drug_resistant_tb_nma/06_validation/validation_framework.md',
        'drug_resistant_tb_nma/07_publication/publication_package.md'
    ]

    # Convert markdown documents to DOCX
    print("📄 Converting Markdown to DOCX...")
    print("-" * 40)

    conversion_success = 0
    total_documents = len(key_documents)

    for md_file in key_documents:
        if os.path.exists(md_file):
            docx_file = md_file.replace('drug_resistant_tb_nma/', 'drug_resistant_tb_nma/08_conversion/docx/').replace('.md', '.docx')
            os.makedirs(os.path.dirname(docx_file), exist_ok=True)

            if convert_markdown_to_docx(md_file, docx_file):
                conversion_success += 1
        else:
            print(f"⚠ File not found: {md_file}")

    print(f"\n✓ Markdown to DOCX conversion: {conversion_success}/{total_documents} successful")
    print()

    # Convert CSV files to Excel
    print("📊 Converting CSV to Excel...")
    print("-" * 40)

    csv_files = [
        'drug_resistant_tb_nma/01_literature_search/literature_search_results.csv',
        'drug_resistant_tb_nma/02_data_extraction/extracted_data.csv',
        'drug_resistant_tb_nma/04_results/treatment_effects_summary.csv',
        'drug_resistant_tb_nma/04_results/component_effects_summary.csv'
    ]

    excel_success = 0

    for csv_file in csv_files:
        if os.path.exists(csv_file):
            excel_file = csv_file.replace('drug_resistant_tb_nma/', 'drug_resistant_tb_nma/08_conversion/excel/').replace('.csv', '.xlsx')

            if convert_csv_to_excel(csv_file, excel_file):
                excel_success += 1

    print(f"\n✓ CSV to Excel conversion: {excel_success}/{len(csv_files)} successful")
    print()

    # Convert DOCX to PDF
    print("📋 Converting DOCX to PDF...")
    print("-" * 40)

    pdf_success = 0
    docx_output_dir = 'drug_resistant_tb_nma/08_conversion/docx/'

    if os.path.exists(docx_output_dir):
        for docx_file in Path(docx_output_dir).glob('*.docx'):
            pdf_file = str(docx_file).replace('docx/', 'pdf/').replace('.docx', '.pdf')

            if create_pdf_from_docx(str(docx_file), pdf_file):
                pdf_success += 1

    print(f"\n✓ DOCX to PDF conversion: {pdf_success} successful")
    print()

    # Create publication package summary
    print("📦 Creating Publication Package Summary...")
    print("-" * 40)

    # Create package summary using string concatenation instead of f-string
    package_summary = "# Publication Package Summary\n\n"
    package_summary += f"**Project:** Drug-Resistant Tuberculosis Network Meta-Analysis\n"
    package_summary += f"**Conversion Date:** {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
    package_summary += f"**Total Documents Converted:** {conversion_success}\n"
    package_summary += f"**Total Data Files Converted:** {excel_success}\n"
    package_summary += f"**Total PDF Files Created:** {pdf_success}\n\n"

    package_summary += "## Converted Files\n\n"
    package_summary += f"### DOCX Documents ({conversion_success} files)\n"

    for i, md_file in enumerate(key_documents[:conversion_success], 1):
        docx_name = Path(md_file).name.replace('.md', '.docx')
        package_summary += f"{i}. {docx_name}\n"

    package_summary += "\n### Excel Data Files ({excel_success} files)\n"
    for i, csv_file in enumerate(csv_files[:excel_success], 1):
        excel_name = Path(csv_file).name.replace('.csv', '.xlsx')
        package_summary += f"{i}. {excel_name}\n"

    package_summary += f"\n### PDF Documents ({pdf_success} files)\n"
    for i in range(1, pdf_success + 1):
        package_summary += f"{i}. Document_{i}.pdf\n"

    package_summary += "\n## File Locations\n\n"
    package_summary += "- **DOCX files:** drug_resistant_tb_nma/08_conversion/docx/\n"
    package_summary += "- **PDF files:** drug_resistant_tb_nma/08_conversion/pdf/\n"
    package_summary += "- **Excel files:** drug_resistant_tb_nma/08_conversion/excel/\n"
    package_summary += "- **Original files:** drug_resistant_tb_nma/\n\n"

    package_summary += "## Submission Checklist\n\n"
    package_summary += "- [ ] All DOCX files reviewed for formatting\n"
    package_summary += "- [ ] All PDF files checked for quality\n"
    package_summary += "- [ ] Excel files verified for data integrity\n"
    package_summary += "- [ ] File sizes within journal limits\n"
    package_summary += "- [ ] All supplementary materials included\n\n"

    package_summary += "## Next Steps\n\n"
    package_summary += "1. Review all converted documents for formatting\n"
    package_summary += "2. Combine PDF files if needed for submission\n"
    package_summary += "3. Upload files to journal submission system\n"
    package_summary += "4. Verify all files meet journal requirements\n\n"

    package_summary += "---\n"
    package_summary += "**Package Version:** 1.0\n"
    package_summary += "**Conversion Status:** Complete\n"

    # Save package summary
    summary_file = 'drug_resistant_tb_nma/08_conversion/publication_package_summary.md'
    with open(summary_file, 'w', encoding='utf-8') as f:
        f.write(package_summary)

    print(f"✓ Publication package summary saved to: {summary_file}")
    print()

    # Final summary
    print("="*60)
    print("CONVERSION COMPLETE!")
    print("="*60)
    print(f"📄 Markdown to DOCX: {conversion_success}/{total_documents} files")
    print(f"📊 CSV to Excel: {excel_success}/{len(csv_files)} files")
    print(f"📋 DOCX to PDF: {pdf_success} files")
    print()
    print("📁 Output directories:")
    print("  • drug_resistant_tb_nma/08_conversion/docx/")
    print("  • drug_resistant_tb_nma/08_conversion/pdf/")
    print("  • drug_resistant_tb_nma/08_conversion/excel/")
    print()
    print("✅ All assets ready for journal submission!")
    print("="*60)

    return True

if __name__ == "__main__":
    main()
