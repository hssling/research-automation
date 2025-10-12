#!/usr/bin/env python3
"""
Simple script to convert Long COVID neurocognitive manuscript from Markdown to DOCX
"""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def convert_md_to_docx(md_file, docx_file):
    """Convert markdown file to DOCX format"""

    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create new document
    doc = Document()

    # Set up document styling
    doc.core_properties.title = "Long COVID Neurocognitive Systematic Review and Meta-Analysis"
    doc.core_properties.author = "Research Automation Platform"

    # Add title
    title = doc.add_heading('Long COVID Neurocognitive Impairment: A Systematic Review and Meta-Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add content (simple approach - split by sections)
    sections = content.split('\n## ')

    for section in sections:
        if not section.strip():
            continue

        lines = section.split('\n')

        # First line is section header
        if lines:
            header_text = lines[0].strip('# ')
            if header_text:
                doc.add_heading(header_text, level=1)

        # Process remaining content
        for line in lines[1:]:
            line = line.strip()
            if not line:
                continue
            elif line.startswith('- '):
                # Bullet point
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(line[2:])
            elif line.startswith('**') and line.endswith('**'):
                # Bold text
                p = doc.add_paragraph()
                run = p.add_run(line.strip('*'))
                run.bold = True
            else:
                # Regular paragraph
                doc.add_paragraph(line)

    # Save document
    doc.save(docx_file)
    print(f"✅ Converted {md_file} to {docx_file}")

if __name__ == "__main__":
    md_file = "long_covid_neurocognitive_manuscript.md"
    docx_file = "long_covid_neurocognitive_manuscript.docx"

    if os.path.exists(md_file):
        convert_md_to_docx(md_file, docx_file)
    else:
        print(f"❌ File not found: {md_file}")
